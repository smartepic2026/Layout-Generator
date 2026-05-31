"""
ingest.py — Notion STEP 3-2 / 3-3 / 3-4 구현
================================================
Notion 가이드의 청킹/임베딩/저장 코드를 그대로 따른다.

* 청킹: 512 토큰(=단어) / 50 오버랩  (Notion STEP 3-2 그대로)
* 메타데이터 스키마: source / doc_type / scale / reliability / year / chunk_index
* 두 컬렉션: design_standards, regulatory_docs
* PDF는 pypdf(=PyPDF2 후속작) 사용. DOCX/XLSX/MD/JSON은 텍스트 추출 후
  동일 청킹 파이프라인에 투입.
"""
from __future__ import annotations
import json
from pathlib import Path
from typing import Dict, Iterable, List

from vector_store import init_db, Collection


# --- Notion STEP 3-2 — chunking -------------------------------------------

def chunk_text(full_text: str, chunk_size: int = 512, overlap: int = 50) -> List[str]:
    """단어(=토큰) 기준 청킹 (Notion 코드 그대로)."""
    if not full_text:
        return []
    words = full_text.split()
    chunks: List[str] = []
    step = max(1, chunk_size - overlap)
    for i in range(0, len(words), step):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk.strip():
            chunks.append(chunk)
        if i + chunk_size >= len(words):
            break
    return chunks


# --- Text extractors (PDF/DOCX/XLSX/MD/JSON) ------------------------------

def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        from PyPDF2 import PdfReader  # type: ignore
    parts = []
    try:
        reader = PdfReader(str(path))
        for p in reader.pages:
            try:
                parts.append(p.extract_text() or "")
            except Exception:
                parts.append("")
    except Exception as e:
        print(f"  ! PDF read error {path.name}: {e}")
    return "\n".join(parts)


def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(path), data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            row_txt = " | ".join("" if v is None else str(v) for v in row)
            if row_txt.strip():
                parts.append(row_txt)
    return "\n".join(parts)


def _extract_md(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_json(path: Path) -> str:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return path.read_text(encoding="utf-8", errors="ignore")
    # flatten to text
    def walk(node, prefix=""):
        out = []
        if isinstance(node, dict):
            for k, v in node.items():
                out.extend(walk(v, f"{prefix}{k}: "))
        elif isinstance(node, list):
            for i, v in enumerate(node):
                out.extend(walk(v, f"{prefix}[{i}] "))
        else:
            out.append(f"{prefix}{node}")
        return out
    return "\n".join(walk(data))


def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext == ".xlsx":
        return _extract_xlsx(path)
    if ext == ".md":
        return _extract_md(path)
    if ext == ".json":
        return _extract_json(path)
    return ""


# --- Notion STEP 3-3 — add documents --------------------------------------

def add_documents_to_collection(collection: Collection, doc_path: Path,
                                metadata: Dict) -> int:
    """Notion 코드 그대로의 시그니처. PDF 외 포맷은 자동 분기."""
    text = extract_text(doc_path)
    chunks = chunk_text(text)
    if not chunks:
        print(f"  ! {doc_path.name}: 0 chunks (skip)")
        return 0
    metadatas = [{**metadata, "chunk_index": i} for i in range(len(chunks))]
    ids = [f"{metadata['source']}_chunk_{i}" for i in range(len(chunks))]
    collection.add(documents=chunks, metadatas=metadatas, ids=ids)
    print(f"  ✅ {metadata['source']}: {len(chunks)}개 청크 저장")
    return len(chunks)


# --- Document inventory (RAG_DB_files 디렉터리 매핑) ----------------------
#
# 분류 원칙:
#   regulatory_docs : 법규/공식 가이던스 (FDA / EU / WHO / ICH / ISO 표준)
#   design_standards: 설계 가이드/실무 예시 (ISPE Baseline, BioPhorum, IPS,
#                     사내 cleanroom 가이드, 의사결정 트리, 장비 레이아웃 예제)
# 분류 근거는 Notion 가이드 Phase 2-5 메타데이터 스키마 + 두 에이전트 역할
# 분리와 동일하다.

REGULATORY = [
    ("1.FDA Guidance Sterile Drug Products 2004.pdf",
        dict(source="FDA_Sterile_Drug_Products_2004", doc_type="regulatory",
             scale="commercial", reliability="high", year=2004,
             jurisdiction="FDA")),
    ("3.EU GMP Annex 1_2022.pdf",
        dict(source="EU_GMP_Annex1_2022", doc_type="regulatory",
             scale="commercial", reliability="high", year=2022,
             jurisdiction="EMA")),
    ("4.WHO TRS Annex-2-gmp-for-sterile-pharmaceutical-products 2022.pdf",
        dict(source="WHO_TRS_Annex2_2022", doc_type="regulatory",
             scale="commercial", reliability="high", year=2022,
             jurisdiction="WHO")),
    ("6.ISO 14644-1_Part1 Classification of air cleanliness.pdf",
        dict(source="ISO_14644_1", doc_type="regulatory",
             scale="commercial", reliability="high", year=2015,
             jurisdiction="ISO")),
    ("7.ISO_14644-4_Part 4 Design, Constructin and Strart-up.docx",
        dict(source="ISO_14644_4", doc_type="regulatory",
             scale="commercial", reliability="high", year=2001,
             jurisdiction="ISO")),
    ("12.FDA Guidance FOR A MONOCLONAL ANTIBODY PRODUCT FOR IN VIVO1996.pdf",
        dict(source="FDA_mAb_Guidance_1996", doc_type="regulatory",
             scale="commercial", reliability="medium", year=1996,
             jurisdiction="FDA")),
]

DESIGN = [
    ("0.GMP_Facility_Layout_References.docx",
        dict(source="GMP_Facility_Layout_References", doc_type="design_standard",
             scale="commercial", reliability="medium", year=2024)),
    ("2.ISPE_Biopharmaceutical facility design-2005.pdf",
        dict(source="ISPE_Biopharm_Facility_2005", doc_type="design_standard",
             scale="commercial", reliability="high", year=2005)),
    ("5.BioPhorum_ BioPhorumImproving the Biomanufacturing Facility.pdf",
        dict(source="BioPhorum_Improving_Biomfg", doc_type="design_standard",
             scale="commercial", reliability="high", year=2021)),
    ("8.Biotech_Cleanroom_Facility_Design_Guide.docx",
        dict(source="Biotech_Cleanroom_Guide", doc_type="design_standard",
             scale="commercial", reliability="medium", year=2024)),
    ("9.Facility_Layout_Design_GMP_Compliance.docx",
        dict(source="Facility_Layout_GMP_Compliance", doc_type="design_standard",
             scale="commercial", reliability="medium", year=2024)),
    ("10.IPS_Biopharmaceutical_Facility_Design_Layout.docx",
        dict(source="IPS_Biopharm_Layout", doc_type="design_standard",
             scale="commercial", reliability="medium", year=2024)),
    ("11.GMP_Gowning_Protocols_Transition_Zones.docx",
        dict(source="GMP_Gowning_TransitionZones", doc_type="design_standard",
             scale="commercial", reliability="medium", year=2024)),
    ("ISPE_Baseline Guide Vol 6 Biopharmaceutical Manufacturing Facilities 3rd Edition_2023 (1).pdf",
        dict(source="ISPE_Baseline_Vol6_2023", doc_type="design_standard",
             scale="commercial", reliability="high", year=2023)),
    ("Equipment_Layout_example.pdf",
        dict(source="Equipment_Layout_Example", doc_type="design_standard",
             scale="commercial", reliability="medium", year=2024)),
    ("GMP Layout Logic_0510.xlsx",
        dict(source="GMP_Layout_Logic_XLSX", doc_type="design_standard",
             scale="commercial", reliability="medium", year=2024)),
    ("GMP_Layout_Decision_Tree.md",
        dict(source="GMP_Layout_DecisionTree_MD", doc_type="design_standard",
             scale="commercial", reliability="medium", year=2024)),
    ("GMP_Layout_Decision_Tree.json",
        dict(source="GMP_Layout_DecisionTree_JSON", doc_type="design_standard",
             scale="commercial", reliability="medium", year=2024)),
]


def run(rag_files_dir: str, db_path: str) -> None:
    src = Path(rag_files_dir)
    client, design, regulatory = init_db(db_path)

    print("\n[regulatory_docs] 인덱싱 시작")
    for fname, meta in REGULATORY:
        p = src / fname
        if not p.exists():
            print(f"  ! 파일 없음: {fname}")
            continue
        add_documents_to_collection(regulatory, p, meta)

    print("\n[design_standards] 인덱싱 시작")
    for fname, meta in DESIGN:
        p = src / fname
        if not p.exists():
            print(f"  ! 파일 없음: {fname}")
            continue
        add_documents_to_collection(design, p, meta)

    # Notion STEP 3-4 — count
    print("\n=== 컬렉션 카운트 ===")
    print(f"design_standards: {design.count():>6}개 청크")
    print(f"regulatory_docs : {regulatory.count():>6}개 청크")

    print("\n[임베딩] TF-IDF (max_features=30000) 학습/적용 중...")
    client.build_embeddings()
    print(f"  vocab size = {len(client.embedder.vocab)}")
    print(f"  design embeddings shape  = {design.embeddings.shape}")
    print(f"  regulatory embeddings shape = {regulatory.embeddings.shape}")

    client.persist()
    print(f"\n저장 완료: {db_path}")


if __name__ == "__main__":
    import sys
    rag_dir = sys.argv[1] if len(sys.argv) > 1 else \
        str(Path(__file__).resolve().parent.parent / "RAG_DB_files")
    db = sys.argv[2] if len(sys.argv) > 2 else \
        str(Path(__file__).resolve().parent / "data")
    run(rag_dir, db)
